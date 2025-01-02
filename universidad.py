from docx import Document

# Crear un nuevo documento de Word para la ficha técnica
doc = Document()

# Agregar el título general
doc.add_heading('Ficha Técnica', level=1)
doc.add_paragraph("________________________________________________________________________________")

# Lista de ejemplos basados en la información extraída
fichas = [
    {
        "Modelo": "Inspiron 5400 AIO",
        "Procesador": "Intel Core i5-1135G7",
        "RAM": "12 GB",
        "Disco Duro": "Mecánico - 1TB - Western Digital",
        "Unidad CD/DVD": "No",
        "Regulador": "No",
        "Monitor": "No",
        "Teclado y Mouse": "Sí - Marca Logisi",
        "Sistema Operativo": "Windows 64 Bits",
        "Ofimática": "Office 2019",
        "Antivirus": "Sí - Vence el 13/01/2024",
        "Estado del Equipo": "Bueno"
    },
    {
        "Modelo": "Inspiron 3480 AIO",
        "Procesador": "Intel Core i5-8265U",
        "RAM": "8 GB",
        "Disco Duro": "Sólido - 220 GB - Kingston",
        "Unidad CD/DVD": "No",
        "Regulador": "No",
        "Monitor": "Sí - Marca Acer",
        "Teclado y Mouse": "Sí - Marca Acteck",
        "Sistema Operativo": "Windows 64 Bits",
        "Ofimática": "Office 2019",
        "Antivirus": "Sí - Vence el 13/01/2024",
        "Estado del Equipo": "Bueno"
    },
    {
        "Modelo": "ThinkStation",
        "Procesador": "Intel Core i5-10600",
        "RAM": "16 GB",
        "Disco Duro": "Mecánico - 1TB - Western Digital",
        "Unidad CD/DVD": "No",
        "Regulador": "No",
        "Monitor": "Sí - Marca Acer",
        "Teclado y Mouse": "Sí - Marca Acteck",
        "Sistema Operativo": "Windows 64 Bits",
        "Ofimática": "Office 2016",
        "Antivirus": "Sí - Vence el 25/03/2024",
        "Estado del Equipo": "Bueno"
    }
]

# Función para agregar una ficha técnica
def agregar_ficha(doc, ficha):
    doc.add_heading(ficha["Modelo"], level=2)
    for key, value in ficha.items():
        if key != "Modelo":
            doc.add_paragraph(f"{key}: {value}")
    doc.add_paragraph("\n________________________________________________________________________________")

# Agregar cada ficha al documento
for ficha in fichas:
    agregar_ficha(doc, ficha)

# Guardar el archivo
output_path = "C:\\Users\\Desarrollo\\Documentos\\Ficha_Tecnica_Transformada.docx"

doc.save(output_path)
output_path
